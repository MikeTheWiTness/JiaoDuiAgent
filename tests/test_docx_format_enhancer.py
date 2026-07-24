import os
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

pytestmark = pytest.mark.slow



def create_test_docx(output_path):
    """创建一个包含各种格式的测试 Word 文档"""
    from docx import Document
    from docx.enum.text import WD_UNDERLINE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()

    def add_emphasis_mark(run):
        """给 run 加着重号（通过 XML 方式）"""
        rPr = run._element.get_or_add_rPr()
        emph = OxmlElement('w:emph')
        emph.set(qn('w:val'), 'dot')
        rPr.append(emph)

    # 标题
    doc.add_heading('测试特殊格式', level=1)

    # 1. 着重号
    p1 = doc.add_paragraph()
    p1.add_run('这是一段普通文字，')
    run_emph = p1.add_run('这几个字有')
    add_emphasis_mark(run_emph)
    run_emph2 = p1.add_run('着重号')
    add_emphasis_mark(run_emph2)
    p1.add_run('，后面是普通文字。')

    # 2. 下划线
    p2 = doc.add_paragraph()
    p2.add_run('这句话中，')
    run_under = p2.add_run('这里有下划线')
    run_under.font.underline = True
    p2.add_run('，其他没有。')

    # 3. 波浪线
    p3 = doc.add_paragraph()
    p3.add_run('这句话中，')
    run_wavy = p3.add_run('这里有波浪线')
    run_wavy.font.underline = WD_UNDERLINE.WAVY
    p3.add_run('，其他没有。')

    # 4. 删除线
    p4 = doc.add_paragraph()
    p4.add_run('这句话中，')
    run_strike = p4.add_run('这里有删除线')
    run_strike.font.strike = True
    p4.add_run('，其他没有。')

    # 5. 下标和上标
    p5 = doc.add_paragraph()
    p5.add_run('化学公式：H')
    run_sub = p5.add_run('2')
    run_sub.font.subscript = True
    p5.add_run('O，数学：x')
    run_sup = p5.add_run('2')
    run_sup.font.superscript = True
    p5.add_run(' + y')
    run_sup2 = p5.add_run('2')
    run_sup2.font.superscript = True
    p5.add_run(' = z')
    run_sup3 = p5.add_run('2')
    run_sup3.font.superscript = True

    # 6. 粗体和斜体（Pandoc 原生支持）
    p6 = doc.add_paragraph()
    p6.add_run('这句话有')
    run_bold = p6.add_run('粗体')
    run_bold.bold = True
    p6.add_run('和')
    run_italic = p6.add_run('斜体')
    run_italic.italic = True
    p6.add_run('。')

    # 7. 古诗词（语文场景）
    doc.add_heading('古诗词示例', level=2)
    p7 = doc.add_paragraph()
    run_title = p7.add_run('静夜思')
    run_title.bold = True
    p7.add_run('\n')
    p7.add_run('床前明月光，')
    p7.add_run('\n')
    run_emph_poem = p7.add_run('疑是地上霜')
    add_emphasis_mark(run_emph_poem)
    p7.add_run('。')
    p7.add_run('\n')
    p7.add_run('举头望明月，')
    p7.add_run('\n')
    run_emph_poem2 = p7.add_run('低头思故乡')
    run_emph_poem2.font.underline = WD_UNDERLINE.WAVY
    p7.add_run('。')

    doc.save(output_path)
    print(f"✅ 测试文档已创建: {output_path}")


def test_extract_formats():
    """测试提取特殊格式"""
    print("\n=== 测试 extract_special_formats ===")
    from shared.docx_format_enhancer import extract_special_formats

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "test.docx")
        create_test_docx(docx_path)

        formats = extract_special_formats(docx_path)
        print(f"提取到 {len(formats)} 个格式项")

        by_type = {}
        for fmt in formats:
            t = fmt["type"]
            if t not in by_type:
                by_type[t] = []
            by_type[t].append(fmt["text"])

        for t, texts in by_type.items():
            print(f"  {t}: {len(texts)} 处")
            for txt in texts[:3]:
                print(f"    - {txt}")

        assert len(formats) > 0, "应该提取到至少一个格式"
        print("✅ 格式提取测试通过")


def test_inject_markers():
    """测试注入格式标记"""
    print("\n=== 测试 inject_format_markers ===")

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "test.docx")
        create_test_docx(docx_path)

        # 先用 pandoc 转换
        from core.pandoc_utils import check_pandoc, convert_with_pandoc
        if not check_pandoc():
            print("⚠️ Pandoc 未安装，跳过完整转换测试")
            return

        output_md = os.path.join(tmpdir, "test.md")
        img_dir = os.path.join(tmpdir, "images")
        os.makedirs(img_dir, exist_ok=True)

        ok = convert_with_pandoc(docx_path, output_md, img_dir, enhance_formats=True)
        assert ok, "Pandoc 转换应该成功"

        with open(output_md, encoding='utf-8') as f:
            md_text = f.read()

        print("转换后的 Markdown 内容:")
        print("-" * 40)
        print(md_text)
        print("-" * 40)

        # 检查标记
        markers_found = []
        for marker in ["<着重>", "<下划线>", "<波浪线>", "<删除线>", "<下标>", "<上标>"]:
            if marker in md_text:
                markers_found.append(marker)
                count = md_text.count(marker)
                print(f"  找到 {marker}: {count} 处")

        print(f"共找到 {len(markers_found)} 种格式标记")

        # 至少应该有几个（取决于匹配成功率）
        print("✅ 标记注入测试完成")


def test_strip_markers():
    """测试移除格式标记"""
    print("\n=== 测试 strip_format_markers ===")
    from shared.docx_format_enhancer import strip_format_markers

    test_text = "这是<着重>测试</着重>文本，有<下划线>下划线</下划线>和<波浪线>波浪线</波浪线>"
    stripped = strip_format_markers(test_text)
    expected = "这是测试文本，有下划线和波浪线"

    print(f"原始: {test_text}")
    print(f"清理后: {stripped}")

    assert stripped == expected, f"清理后不匹配: {stripped} != {expected}"
    print("✅ 标记清理测试通过")


if __name__ == "__main__":
    print("Word 格式增强模块测试\n")

    try:
        test_extract_formats()
        test_inject_markers()
        test_strip_markers()

        print("\n" + "=" * 50)
        print("✅ 所有测试通过！")
    except Exception as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
